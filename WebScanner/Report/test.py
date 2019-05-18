#--Created by WD
#python 3.6
#coding:utf-8


#coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

document = Document()

paragraph_font_name = u'宋体'
title_font_name = u'黑体'

#添加标题
title = document.add_heading(u'网站漏洞扫描报告', level = 0)
title_run = title.add_run()
title_run.font.size = Pt(24)
title_run.font.name = title_font_name
title_r = title_run._element
title_r.rPr.rFonts.set(qn('w:eastAsia'), title_font_name)


paragraph = document.add_paragraph()
run = paragraph.add_run(u'中文内容')
run.font.size = Pt(24)
run.font.name = paragraph_font_name
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), paragraph_font_name)

document.save('111.docx')