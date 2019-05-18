#--Created by WD
#python 3.6
#coding:utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#设置文档默认字体
document.styles['Normal'].font.name = u'微软雅黑'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
title = document.add_paragraph()
#大标题居中
title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

run = document.add_heading('', level=3)
run.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = run.add_run(u"应用场景示例: ")  # 应用场景示例标题
run.font.name = u'黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

run = document.add_heading('', level=1).add_run(u"应用场景示例: ")  # 应用场景示例标题
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


# 参数#文档对象，文字内容，文字大小，文字样式（目前就只判断了粗体）
def writeP(document, content, size, style = None):
    p = document.add_paragraph()
    run = p.add_run(content)
    font = run.font
    font.size = Pt(size)
    if style == 'bold': font.bold = True

document.save('22.docx')