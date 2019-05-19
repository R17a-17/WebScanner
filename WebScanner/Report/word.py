#--Created by WD
#python 3.6
#coding:utf-8

#######################本脚本用于处理word文档，生成扫描报告###############################


###############################################
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
import re
from urllib.parse import urlparse
##################################################
#内部类
from WebScanner.Report import histogram_pic
from WebScanner.Report import piechart_pic
from WebScanner.Mysqldb import GetVulnAPI
##################################################


#标题字体
title_font_name = u'黑体'
#标题大小22，为二号大小（24为小一，26为一号，14为四号）
top_title_size = 22
sec_title_size = 14

#图片宽度为6英寸
picture_width = 6

class GenerateReport(object):
    '''生成漏洞扫描报告'''

    def __init__(self,sitename,vulnnum, urgentvulnnum, highvulnnum, midumvulnnum, sitedomain, scantime, scanurlnum,piechartdata,vulninfo,fname = 'report.docx'):
        # 生成一个word文档对象
        self.document = Document()
        # 设置文档默认字体
        self.document.styles['Normal'].font.name = u'宋体'
        # 设置默认字体大小 小四
        self.document.styles['Normal'].font.size = Pt(12)
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        self.setTitle(sitename)
        self.setSummaryPart(sitename,vulnnum, urgentvulnnum, highvulnnum, midumvulnnum, sitedomain, scantime, scanurlnum)
        self.setGraph((0,0,midumvulnnum,highvulnnum,urgentvulnnum),piechartdata)
        self.setVulnlist(vulnnum,vulninfo)

        # path = '../reporttmp/'
        # self.document.save(path + fname)


    def setTitle(self, titlestr):
        # 添加标题：参数为标题字符串
        titlestr = '"' + titlestr + '"网站安全报告'
        title = self.document.add_heading('', level=1)
        title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 标题居中
        title_run = title.add_run(titlestr)
        title_run.font.size = Pt(top_title_size)  # 设置大小
        title_run.font.name = title_font_name  # 设置字体
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font_name)


    def setSummaryPart(self,sitename,vulnnum, urgentvulnnum, highvulnnum, midumvulnnum, sitedomain, scantime, scanurlnum):
        '''设置概况部分'''
        # 一、网站概况
        summary_titlestr = '一、网站概况'
        summary_title = self.document.add_heading('', level=2)
        summary_titele_run = summary_title.add_run(summary_titlestr)
        summary_titele_run.font.name = u'宋体'
        summary_titele_run.bold = True  # 加粗
        summary_titele_run.font.size = Pt(sec_title_size)
        summary_titele_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        vulnnum = str(vulnnum)
        highvulnnum = str(highvulnnum)
        urgentvulnnum = str(urgentvulnnum)
        midumvulnnum = str(midumvulnnum)
        summary_para = '    本次针对' + sitename + '网站进行扫描，共发现' + vulnnum + '个漏洞，其中紧急漏洞' + urgentvulnnum + '个，' \
                                                                                                       '高危漏洞' + highvulnnum + '个，中危漏洞' + midumvulnnum + '个，扫描共用时长为' \
                                                                                                        + scantime + '。网站基本信息如下：'
        self.document.add_paragraph(summary_para)

        # 基本信息列表
        sitedomain = sitedomain
        scanurlnum = str(scanurlnum)
        scantime = scantime
        summarylist_domain = '域名：' + sitedomain
        summarylist_urlnum = '扫描url总数：' + scanurlnum
        summarylist_time = '扫描时长：' + scantime
        summarylist = [summarylist_domain, summarylist_urlnum, summarylist_time]
        i = 0
        while i < 3:
            self.document.add_paragraph(
                summarylist[i], style='ListNumber'
            )
            i = i + 1


    def setGraph(self,histogramdata,piechartdata):
        '''漏洞分布情况展示'''
        titlestr = '二、漏洞分布情况'
        title = self.document.add_heading('', level=2)
        titele_run = title.add_run(titlestr)
        titele_run.font.name = u'宋体'
        titele_run.bold = True  # 加粗
        titele_run.font.size = Pt(sec_title_size)
        titele_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        #添加柱状图
        histogram_picname = histogram_pic.make_histogrampng(histogramdata)
        self.document.add_picture(histogram_picname, width=Inches(picture_width))
        para = self.document.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph = para.add_run(u'等级漏洞情况')
        paragraph.font.size = Pt(10)
        paragraph.font.name = u'微软雅黑'
        paragraph._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        #添加饼图
        piechart_picname = piechart_pic.make_piechartpng(piechartdata)
        self.document.add_picture(piechart_picname, width=Inches(picture_width))
        para1 = self.document.add_paragraph()
        para1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph = para1.add_run(u'各类型漏洞分布')
        paragraph.font.size = Pt(10)
        paragraph.font.name = u'微软雅黑'
        paragraph._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')


    def setVulnlist(self,vulnnum,vulninfo):
        '''生成漏洞详情列表'''
        titlestr = '三、漏洞详情列表'
        title = self.document.add_heading('', level=2)
        titele_run = title.add_run(titlestr)
        titele_run.font.name = u'宋体'
        titele_run.bold = True  # 加粗
        titele_run.font.size = Pt(sec_title_size)
        titele_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        #生成列表
        # 漏洞类型： XSS
        # 漏洞等级：高危
        # 漏洞url：http://xxx/UploadFile/picture/
        # 漏洞危害：XXXXXXXXXXXXX
        # 修复建议：XXXXXXXXXXX

        while vulnnum > 0:
            self.document.add_paragraph('                     \n')
            table = self.document.add_table(rows=5, cols=2)
            hdr_cells1 = table.rows[0].cells
            hdr_cells1[0].text = '漏洞类型'
            hdr_cells1[1].text = vulninfo[vulnnum-1]['vulntype']
            hdr_cells2 = table.rows[1].cells
            hdr_cells2[0].text = '漏洞等级'
            hdr_cells2[1].text = vulninfo[vulnnum-1]['vulnlevel']
            hdr_cells3 = table.rows[2].cells
            hdr_cells3[0].text = '漏洞URL'
            hdr_cells3[1].text = vulninfo[vulnnum-1]['vulnurl']
            hdr_cells4 = table.rows[3].cells
            hdr_cells4[0].text = '漏洞危害'
            hdr_cells4[1].text = vulninfo[vulnnum-1]['vulnaffection']
            hdr_cells5 = table.rows[4].cells
            hdr_cells5[0].text = '修复建议'
            hdr_cells5[1].text = vulninfo[vulnnum-1]['vulnsuggestion']

            vulnnum = vulnnum - 1


    def savefile(self,fname):
        self.document.save(fname)


def get_title(url):
    '''
    用re抽取网页Title
    '''

    html = requests.get(url).text
    compile_rule = r'<title>.*</title>'
    title_list = re.findall(compile_rule, html)
    if title_list == []:
        title = ''
    else:
        title = title_list[0][7:-8]
    return title


def main(url,scantime,scanurlnum):
    sitename = get_title(url)
    sql = GetVulnAPI.GetVuln()
    vulnnum = sql.getAllVuln_num()
    Weakpwdnum = sql.getWeakpwdVuln_num()
    Xssnum = sql.getXssVuln_num()
    Sqlinum = sql.getSqliVuln_num()
    Crlfnum = sql.getCrlfVuln_num()
    highvulnnum = Weakpwdnum + Xssnum
    urgentvulnnum = Sqlinum
    midumvulnnum = Crlfnum

    sitedomain = urlparse(url).netloc
    piechartdata = [Xssnum,Sqlinum,Crlfnum,Weakpwdnum]

    vulninfo = sql.getAllVuln_url()

    doc = GenerateReport(
        sitename=sitename, vulnnum=vulnnum, urgentvulnnum=urgentvulnnum, highvulnnum=highvulnnum, midumvulnnum=midumvulnnum,
        sitedomain=sitedomain, scantime=scantime, scanurlnum=scanurlnum,
        piechartdata=piechartdata,
        vulninfo = vulninfo,
        )
    return doc



if __name__ == '__main__':
    word = main(url = 'http://192.168.177.161/bWAPP_latest/bwapp/login.php',scantime='5分10秒',scanurlnum=50)

